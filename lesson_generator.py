import os
import io
import streamlit as st
from google import genai
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from fpdf import FPDF

# Load environment variables
load_dotenv()

# --- Page Configuration ---
st.set_page_config(page_title="Personalized Lesson Creator", layout="wide")

# --- AI Client Setup ---
gemini_key = os.getenv("GEMINI_API_KEY")

def get_ai_client():
    if not gemini_key:
        st.error("Missing GEMINI_API_KEY. Please set it in your environment.")
        return None
    return genai.Client(api_key=gemini_key)

AI_MODEL = 'gemini-2.0-flash'

# --- Helper Functions for File Conversion ---

def clean_for_pdf(text):
    """Replaces curly quotes and dashes that crash PDF generators."""
    replacements = {
        '\u2018': "'", '\u2019': "'",  # Curly single
        '\u201c': '"', '\u201d': '"',  # Curly double
        '\u2013': '-', '\u2014': '-',  # Dashes
        '\u2022': '-',                # Replace dots with hyphens for safety
        '*': '-'                      # Convert any AI stars to clean hyphens
    }
    for unicode_char, safe_char in replacements.items():
        text = text.replace(unicode_char, safe_char)
    return text.encode('latin-1', 'ignore').decode('latin-1')

def create_docx(text):
    """Generates a professional Word document with clean bullet points."""
    doc = Document()
    doc.add_heading('Tailored Lesson Plan', 0)
    
    for line in text.split('\n'):
        clean_line = line.strip()
        if not clean_line:
            continue
            
        headers = ["Objective:", "Concept Breakdown:", "Activity:", "Inclusivity Considerations:", "Wrap Up:"]
        
        # If it's a header
        if any(head in clean_line for head in headers):
            run = doc.add_paragraph().add_run(clean_line)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(20, 80, 150)
        # If it's a bullet point (starts with -)
        elif clean_line.startswith('-'):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(clean_line.lstrip('- '))
        else:
            doc.add_paragraph(clean_line)
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_pdf(text):
    """Generates a professional PDF using hyphens for bullet points."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    pdf.set_font("Helvetica", 'B', 16)
    pdf.cell(0, 10, "Tailored Lesson Plan", ln=True, align='C')
    pdf.ln(5)

    safe_text = clean_for_pdf(text)

    for line in safe_text.split('\n'):
        clean_line = line.strip()
        if not clean_line:
            pdf.ln(2)
            continue
            
        headers = ["Objective:", "Concept Breakdown:", "Activity:", "Inclusivity Considerations:", "Wrap Up:"]
        if any(head in clean_line for head in headers):
            pdf.set_font("Helvetica", 'B', 12)
            pdf.set_text_color(20, 80, 150)
            pdf.multi_cell(0, 10, clean_line)
            pdf.set_text_color(0, 0, 0)
        elif clean_line.startswith('-'):
            pdf.set_font("Helvetica", size=11)
            # Indent bullet points slightly
            pdf.set_x(15)
            pdf.multi_cell(0, 8, clean_line)
        else:
            pdf.set_font("Helvetica", size=11)
            pdf.multi_cell(0, 8, clean_line)
            
    return bytes(pdf.output())

# --- UI Layout ---
st.title("🍎 AI Lesson Personalizer")

if "tailored_lesson" not in st.session_state:
    st.session_state.tailored_lesson = ""

col1, col2 = st.columns([1, 1], gap="medium")

with col1:
    st.subheader("Input Details")
    student_interests = st.text_input("Student Interests", placeholder="e.g. NASA, Minecraft")
    base_lesson = st.text_area("Lesson Content", height=300)
    
    if st.button("Generate Lesson with Bullets", type="primary"):
        if not student_interests or not base_lesson:
            st.warning("Please fill in both fields.")
        else:
            client = get_ai_client()
            if client:
                with st.spinner("Writing..."):
                    prompt = f"""
                    You are an expert inclusive education specialist.
                    Interest: {student_interests}
                    Context: {base_lesson}
                    
                    TASK: Rewrite as a clean 5-minute lesson.
                    
                    FORMATTING RULES:
                    1. Use exactly these headers: "Objective:", "Concept Breakdown:", "Activity:", "Inclusivity Considerations:", and "Wrap Up:".
                    2. Use hyphens (-) for bullet points. 
                    3. DO NOT use asterisks (*) anywhere.
                    4. Keep sentences clear and concise.
                    """
                    try:
                        response = client.models.generate_content(model=AI_MODEL, contents=prompt)
                        # Swap any accidental asterisks with hyphens for the UI
                        st.session_state.tailored_lesson = response.text.replace('*', '-')
                    except Exception as e:
                        st.error(f"Error: {e}")

with col2:
    st.subheader("Tailored Output")
    if st.session_state.tailored_lesson:
        st.write(st.session_state.tailored_lesson)
        
        st.divider()
        dcol1, dcol2 = st.columns(2)
        with dcol1:
            st.download_button("Download Word", data=create_docx(st.session_state.tailored_lesson), file_name="lesson.docx")
        with dcol2:
            st.download_button("Download PDF", data=create_pdf(st.session_state.tailored_lesson), file_name="lesson.pdf")